# modules/utils.py
import PyPDF2
from docx import Document
from pathlib import Path

def extract_text_from_pdf(file):
    """
    Extracts text from a PDF file object or path.
    """
    if isinstance(file, (str, Path)):
        reader = PyPDF2.PdfReader(str(file))
    else:
        reader = PyPDF2.PdfReader(file)
    
    text = ""
    for page in reader.pages:
        text += page.extract_text() + " "
    return text

def extract_text_from_docx(file):
    """
    Extracts text from a DOCX file object or path.
    """
    if isinstance(file, (str, Path)):
        doc = Document(str(file))
    else:
        doc = Document(file)
    
    text = ""
    for para in doc.paragraphs:
        text += para.text + " "
    return text

def extract_skills_from_resume(file):
    """
    Extracts capitalized words from resume as a basic skill list.
    Accepts file object, file path, or string.
    """
    text = ""
    if hasattr(file, "name"):
        # File object
        if file.name.endswith(".pdf"):
            text = extract_text_from_pdf(file)
        elif file.name.endswith(".docx"):
            text = extract_text_from_docx(file)
        else:
            text = str(file.read(), "utf-8")
    elif isinstance(file, (str, Path)):
        # File path
        path = Path(file)
        if path.suffix == ".pdf":
            text = extract_text_from_pdf(path)
        elif path.suffix == ".docx":
            text = extract_text_from_docx(path)
        else:
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()
    else:
        # Plain string
        text = str(file)

    # Extract words starting with uppercase letters as skills
    skills = [word.strip() for word in text.split() if word.istitle()]
    return skills